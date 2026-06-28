import os
import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def buscar_cnpj(cnpj):
    """Busca dados públicos do CNPJ usando a API pública do ReceitaWS."""
    cnpj_limpo = "".join(filter(str.isdigit, cnpj))
    print(f"[*] Iniciando varredura técnica no CNPJ: {cnpj_limpo}...")
    url = f"https://receitaws.com.br/v1/cnpj/{cnpj_limpo}"
    
    try:
        response = requests.get(url, timeout=15)
        if response.status_code == 200:
            dados = response.json()
            if dados.get("status") == "ERROR":
                print(f"[-] Erro na API: {dados.get('message')}")
                return None
            return dados
        else:
            print(f"[-] Falha na requisição. Código HTTP: {response.status_code}")
            return None
    except Exception as e:
        print(f"[-] Erro de conexão: {str(e)}")
        return None

def set_cell_background(cell, fill_hex):
    """Aplica cor de fundo a uma célula de tabela no Word."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def gerar_relatorio_docx(dados, filename="Relatorio_Inteligencia_CNPJ.docx"):
    """Gera um relatório profissional altamente formatado (.docx)."""
    doc = Document()
    
    # Configurações de Cores
    CHAVE_COR = RGBColor(15, 23, 42)    # Azul Escuro / Cinza Escuro Técnico
    CINZA_TEXTO = RGBColor(71, 85, 105)
    
    # Cabeçalho Principal
    title = doc.add_paragraph()
    title_run = title.add_run("RELATÓRIO DE INTELIGÊNCIA CORPORATIVA (OSINT)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = CHAVE_COR
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Alvo Técnico: {dados.get('nome')} | CNPJ: {dados.get('cnpj')}")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = CINZA_TEXTO
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n")
    
    # Seção 1: Dados Cadastrais Estruturais
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Informações Cadastrais Principais")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(13)
    h1_run.font.bold = True
    h1_run.font.color.rgb = CHAVE_COR
    
    # Tabela de Dados
    info_dict = {
        "Razão Social": dados.get("nome"),
        "Nome Fantasia": dados.get("fantasia") or "Não Informado",
        "Abertura": dados.get("abertura"),
        "Situação Cadastral": dados.get("situacao"),
        "Natureza Jurídica": dados.get("natureza_juridica"),
        "Capital Social": f"R$ {dados.get('capital_social'):,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
        "Logradouro": f"{dados.get('logradouro')}, Nº {dados.get('numero')} {dados.get('complemento')}",
        "Bairro/CEP": f"{dados.get('bairro')} | CEP: {dados.get('cep')}",
        "Município/UF": f"{dados.get('municipio')} - {dados.get('uf')}",
        "Telefone/E-mail": f"{dados.get('telefone')} | {dados.get('email') or 'Não Informado'}"
    }
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Shading Accent 1'
    
    for campo, valor in info_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = campo
        row_cells[1].text = str(valor)
        # Formatando texto interno
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_background(row_cells[0], "F1F5F9")
        
    doc.add_paragraph("\n")
    
    # Seção 2: Quadro de Sócios e Administradores (QSA)
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Quadro de Sócios e Administradores (QSA)")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2_run.font.color.rgb = CHAVE_COR
    
    qsa = dados.get("qsa", [])
    if qsa:
        table_qsa = doc.add_table(rows=1, cols=2)
        table_qsa.style = 'Light Shading Accent 1'
        hdr_cells = table_qsa.rows[0].cells
        hdr_cells[0].text = 'Nome do Sócio/Administrador'
        hdr_cells[1].text = 'Qualificação / Cargo'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        set_cell_background(hdr_cells[0], "CBD5E1")
        set_cell_background(hdr_cells[1], "CBD5E1")
        
        for socio in qsa:
            row_cells = table_qsa.add_row().cells
            row_cells[0].text = socio.get("nome")
            row_cells[1].text = socio.get("qual")
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph("Nenhum sócio ou administrador listado nas bases públicas primárias.")
        
    doc.add_paragraph("\n")
    
    # Seção 3: Atividades Econômicas
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Atividades Econômicas Cadastradas")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(13)
    h3_run.font.bold = True
    h3_run.font.color.rgb = CHAVE_COR
    
    doc.add_paragraph("Atividade Principal:").runs[0].font.bold = True
    for p in dados.get("atividade_principal", []):
        doc.add_paragraph(f"• {p.get('code')} - {p.get('text')}")
        
    doc.add_paragraph("\nAtividades Secundárias:").runs[0].font.bold = True
    for s in dados.get("atividades_secundarias", []):
        doc.add_paragraph(f"• {s.get('code')} - {s.get('text')}")
        
    doc.save(filename)
    print(f"[+] Relatório executivo salvo com sucesso: '{filename}'")

def gerar_painel_html(cnpj, dados):
    """Gera uma página HTML tática com links OSINT dinâmicos e clicáveis."""
    cnpj_limpo = "".join(filter(str.isdigit, cnpj))
    razao_social = dados.get("nome", "Alvo Desconhecido")
    
    html_template = f"""<!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <title>Painel OSINT de Investigação - {cnpj_limpo}</title>
        <style>
            body {{ font-family: 'Segoe UI', Arial, sans-serif; background-color: #0f172a; color: #f8fafc; margin: 0; padding: 25px; }}
            .container {{ max-width: 950px; margin: 0 auto; background-color: #1e293b; padding: 35px; border-radius: 12px; box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.5); border-left: 6px solid #ef4444; }}
            h1 {{ color: #ef4444; margin-top: 0; font-size: 26px; text-transform: uppercase; border-bottom: 2px solid #334155; padding-bottom: 12px; }}
            .meta {{ background: #0f172a; padding: 15px; border-radius: 8px; margin-bottom: 30px; border: 1px solid #334155; }}
            .meta p {{ margin: 5px 0; color: #cbd5e1; font-size: 15px; }}
            .grid {{ display: table; width: 100%; border-collapse: separate; border-spacing: 20px 0px; }}
            .col {{ display: table-cell; width: 50%; vertical-align: top; }}
            .card {{ background: #111827; padding: 22px; border-radius: 8px; border: 1px solid #374151; margin-bottom: 20px; }}
            .card h3 {{ color: #38bdf8; margin-top: 0; font-size: 18px; border-bottom: 1px solid #374151; padding-bottom: 8px; }}
            .btn {{ display: block; width: 100%; padding: 12px; margin: 12px 0; background: #1f2937; color: #f3f4f6; text-align: center; text-decoration: none; border-radius: 6px; font-weight: bold; border: 1px solid #4b5563; transition: all 0.2s; box-sizing: border-box; }}
            .btn:hover {{ background: #ef4444; color: #fff; border-color: #ef4444; transform: translateY(-1px); }}
            .footer {{ text-align: center; margin-top: 40px; font-size: 12px; color: #64748b; letter-spacing: 0.5px; }}
        </style>
    </head>
    <body>

    <div class="container">
        <h1>Painel de Alvos e Fontes Abertas (OSINT)</h1>
        
        <div class="meta">
            <p><strong>Razão Social Alvo:</strong> {razao_social}</p>
            <p><strong>CNPJ Analisado:</strong> {dados.get('cnpj')} (Apenas dígitos: {cnpj_limpo})</p>
            <p><strong>Status Cadastral:</strong> {dados.get('situacao')}</p>
        </div>

        <div class="grid">
            <div class="col">
                <div class="card">
                    <h3>Análise de Processos e Litígios</h3>
                    <p style="font-size:12px; color:#94a3b8;">Verificação de fraudes, disputas civis, criminais ou trabalhistas:</p>
                    <a href="https://www.jusbrasil.com.br/busca?q={cnpj_limpo}" target="_blank" class="btn">Pesquisar no Jusbrasil</a>
                    <a href="https://esaj.tjsp.jus.br/cpopg/search.do?dadosConsulta.valorConsulta={cnpj_limpo}" target="_blank" class="btn">Pesquisar no Tribunal (e-SAJ)</a>
                    <a href="https://www.jusbrasil.com.br/diarios/busca?q={cnpj_limpo}" target="_blank" class="btn">Buscar em Diários Oficiais</a>
                </div>
            </div>
            
            <div class="col">
                <div class="card">
                    <h3>Rastreamento e Pegada Digital</h3>
                    <p style="font-size:12px; color:#94a3b8;">Doxxing avançado, contratos públicos e vazamentos na internet:</p>
                    <a href="https://www.portaltransparencia.gov.br/busca?termo={cnpj_limpo}" target="_blank" class="btn">Portal da Transparência Federal</a>
                    <a href="https://www.google.com/search?q=%22{dados.get('cnpj')}%22+OR+%22{cnpj_limpo}%22" target="_blank" class="btn">Doxxing Google (Termo Exato)</a>
                    <a href="https://consultacnpj.com/cnpj/{cnpj_limpo}" target="_blank" class="btn">Dossiê Espelho Alternativo</a>
                </div>
            </div>
        </div>

        <div class="footer">
            Operação de Segurança da Informação e Inteligência Forense Digital • Will Oliveira Dev
        </div>
    </div>

    </body>
    </html>
    """
    
    html_filename = f"painel_osint_{cnpj_limpo}.html"
    with open(html_filename, "w", encoding="utf-8") as f:
        f.write(html_template)
    print(f"[+] Painel tático HTML gerado com sucesso: '{html_filename}'")

# ==========================================
# EXECUÇÃO DO FLUXO OPERACIONAL
# ==========================================
if __name__ == "__main__":
    # Exemplo de CNPJ público ativo para teste (Substitua pelo CNPJ do seu alvo)
    CNPJ_ALVO = "27.865.757/0001-02" 
    
    # 1. Executa a varredura
    dados_obtidos = buscar_cnpj(CNPJ_ALVO)
    
    if dados_obtidos:
        cnpj_numeros = "".join(filter(str.isdigit, CNPJ_ALVO))
        
        # 2. Gera os artefatos técnicos de saída
        gerar_relatorio_docx(dados_obtidos, filename=f"Relatorio_OSINT_{cnpj_numeros}.docx")
        gerar_painel_html(CNPJ_ALVO, dados_obtidos)
        
        print("\n[✔] Missão concluída. Arquivos prontos para análise.")
    else:
        print("\n[-] Falha crítica na coleta. Operação abortada.")
