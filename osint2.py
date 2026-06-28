# -*- coding: utf-8 -*-
"""
================================================================================
           OSINT - INTELIGÊNCIA FORENSE
                   WILLIAN DE OLIVEIRA
================================================================================
Engenharia de Software Aplicada à Investigação Digital e Inteligência Cibernética
Unificação de Módulos Cadastrais, Financeiros, Judiciais e Análise de Telegram
================================================================================
"""

import os
import re
import sys
import asyncio
import requests
import webbrowser
import urllib.parse
from datetime import datetime

# Importação da biblioteca para relatórios em texto editável
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURAÇÕES DE INTERFACE DO TERMINAL (ANSI COLORS) ---
GREEN  = "\033[92m"
RED    = "\033[91m"
BLUE   = "\033[94m"
CYAN   = "\033[96m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

BANNER_PRINCIPAL = f"""{CYAN}{BOLD}
===========================================================================
           INTELIGÊNCIA OSINT
===========================================================================
     Operador: Willian de Oliveira
==========================================================================={RESET}"""


class WillianOsintSupremo:

    def __init__(self):
        self.dirs = {
            "civil": "DOSSIES_OSINT",
            "juridico": "DOSSIES_JURIDICOS",
            "financeiro": "RELATORIOS_FINANCEIROS",
            "centralizado": "INVESTIGACOES_WILLIAN"
        }
        self.inicializar_pastas()
        self.tg_api_id = "37275908"
        self.tg_api_hash = "cd5599096bbd2de1763339c25de37676"

    def inicializar_pastas(self):
        for pasta in self.dirs.values():
            if not os.path.exists(pasta):
                os.makedirs(pasta)

    def limpar_tela(self):
        os.system('cls' if os.name == 'nt' else 'clear')

    def limpar_documento(self, doc):
        return re.sub(r'[^0-9]', '', doc)

    def obter_info_ddd(self, phone_number):
        if len(phone_number) < 2:
            return "Localidade: Indeterminada"
        ddd = phone_number[:2]
        mapeamento_ddd = {
            "11": "São Paulo (Capital e Região Metropolitana)",
            "12": "São Paulo (Vale do Paraíba e Litoral Norte)",
            "13": "São Paulo (Baixada Santista e Vale do Ribeira)",
            "14": "São Paulo (Bauru, Marília, Jaú, Botucatu)",
            "15": "São Paulo (Sorocaba, Itapetininga)",
            "16": "São Paulo (Ribeirão Preto, Franca, Araraquara, São Carlos)",
            "17": "São Paulo (São José do Rio Preto, Araçatuba, Barretos)",
            "18": "São Paulo (Presidente Prudente, Araçatuba, Assis)",
            "19": "São Paulo (Campinas, Piracicaba, Limeira, Americana)",
            "21": "Rio de Janeiro (Capital e Região Metropolitana)",
            "22": "Rio de Janeiro (Norte Fluminense, Região dos Lagos e Serrana)",
            "24": "Rio de Janeiro (Regiao Sul, Serrana e Costa Verde)",
            "27": "Espírito Santo (Vitória e Região Metropolitana)",
            "31": "Minas Gerais (Belo Horizonte, Região Metropolitana)",
            "34": "Minas Gerais (Uberlândia, Uberaba e Triângulo Mineiro)",
            "41": "Paraná (Curitiba e Região Metropolitana)",
            "47": "Santa Catarina (Joinville, Blumenau, Balneário Camboriú)",
            "48": "Santa Catarina (Florianópolis, Criciúma e Região Sul)",
            "51": "Rio Grande do Sul (Porto Alegre e Região Metropolitana)",
            "61": "Distrito Federal e Entorno",
            "62": "Goiás (Goiânia, Região Metropolitana e Anápolis)",
            "67": "Mato Grosso do Sul (Campo Grande)",
            "71": "Bahia (Salvador e Região Metropolitana)",
            "81": "Pernambuco (Recife e Região Metropolitana)",
            "85": "Ceará (Fortaleza e Região Metropolitana)",
            "91": "Pará (Belém, Região Metropolitana)",
            "92": "Amazonas (Manaus, Região Metropolitana)"
        }
        return mapeamento_ddd.get(ddd, f"DDD {ddd} - Estado identificado via index")

    # =========================================================================
    # MOTOR DE GERAÇÃO EM WORD (.DOCX) - PADRÃO EXECUTIVO
    # =========================================================================
    def gerar_backup_word(self, pasta_destino, nome_base, titulo_doc, qualificacao, links):
        """Gera uma versão em Word paralela e limpa para edição posterior e inserção de prints"""
        try:
            doc = Document()
            cor_principal = RGBColor(26, 43, 76)  # Azul Escuro Executivo
            cor_texto = RGBColor(44, 62, 80)     # Cinza Leitura
            
            # Título Principal
            t = doc.add_paragraph()
            t_run = t.add_run(titulo_doc)
            t_run.font.name = 'Arial'
            t_run.font.size = Pt(18)
            t_run.font.bold = True
            t_run.font.color.rgb = cor_principal
            
            # Metadados
            meta = doc.add_paragraph()
            m_run = meta.add_run(f"Data de Emissão: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\nResponsável: Inteligência Tecnológica\nStatus: Pronto para Inserção de Evidências")
            m_run.font.name = 'Arial'
            m_run.font.size = Pt(10)
            m_run.font.italic = True
            m_run.font.color.rgb = RGBColor(127, 140, 141)
            
            doc.add_paragraph("_" * 50)
            
            # Seção I - Qualificação
            h1 = doc.add_paragraph()
            h1_run = h1.add_run("I - Detalhes e Qualificação do Alvo")
            h1_run.font.name = 'Arial'
            h1_run.font.size = Pt(13)
            h1_run.font.bold = True
            h1_run.font.color.rgb = cor_principal
            
            for k, v in qualificacao.items():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r_key = p.add_run(f"• {k}: ")
                r_key.bold = True
                r_key.font.name = 'Arial'
                p.add_run(str(v)).font.name = 'Arial'
            
            # Espaço demarcado para Prints Gerais
            doc.add_paragraph("\n")
            h_print1 = doc.add_paragraph()
            hp1_run = h_print1.add_run("[ INSIRA AQUI OS PRINTS DE QUALIFICAÇÃO / PAINEL GERAL ]")
            hp1_run.font.name = 'Arial'
            hp1_run.font.size = Pt(11)
            hp1_run.font.bold = True
            hp1_run.font.color.rgb = RGBColor(192, 57, 43) # Vermelho Alerta
            doc.add_paragraph("(Clique nesta linha, apague este texto e cole o print correspondente)\n")

            # Seção II - Vetores de Busca
            h2 = doc.add_paragraph()
            h2_run = h2.add_run("II - Fontes de Cruzamento Vinculadas (Vetores Auditados)")
            h2_run.font.name = 'Arial'
            h2_run.font.size = Pt(13)
            h2_run.font.bold = True
            h2_run.font.color.rgb = cor_principal
            
            for k, v in links.items():
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                r_link_title = p.add_run(f"• {k}:\n")
                r_link_title.bold = True
                r_link_url = p.add_run(v)
                r_link_url.font.color.rgb = RGBColor(41, 128, 185)
                r_link_url.font.underline = True
            
            # Espaço demarcado para Prints Técnicos/DNS
            doc.add_paragraph("\n")
            h_print2 = doc.add_paragraph()
            hp2_run = h_print2.add_run("[ INSIRA AQUI OS PRINTS DE INFRAESTRUTURA / PROVAS ADICIONAIS ]")
            hp2_run.font.name = 'Arial'
            hp2_run.font.size = Pt(11)
            hp2_run.font.bold = True
            hp2_run.font.color.rgb = RGBColor(192, 57, 43)
            doc.add_paragraph("(Clique nesta linha, apague este texto e cole o seu print técnico final)\n")
            
            # Salvamento do arquivo Word (.docx)
            word_filename = f"{nome_base}.docx"
            word_filepath = os.path.join(pasta_destino, word_filename)
            doc.save(word_filepath)
            print(f"{GREEN}[📝] ARQUIVO WORD (.DOCX) COMPATÍVEL GERADO EM: {word_filepath}{RESET}")
        except Exception as e:
            print(f"{RED}[-] Falha ao gerar arquivo estruturado do Word: {e}{RESET}")

    # =========================================================================
    # ESTEIRA 1: OSINT CIVIL / INFRAESTRUTURA
    # =========================================================================
    def criar_dossie_html_civil(self, categoria, alvo, dados_extra, links_gerados):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        timestamp_humano = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        alvo_limpo = re.sub(r'[^a-zA-Z0-9]', '_', alvo)
        
        nome_base_arquivo = f"DOSSIE_{categoria}_{alvo_limpo}_{timestamp}"
        filename_html = f"{nome_base_arquivo}.html"
        filepath_html = os.path.join(self.dirs["civil"], filename_html)
        
        linhas_tabela = "".join([f'<tr><td class="label">{k}</td><td class="value">{v}</td></tr>' for k, v in dados_extra.items()])
        botoes_links = "".join([f'<div class="link-card"><span class="link-title">{k}</span><a href="{v}" target="_blank" class="btn-link">Disparar Busca Passiva</a></div>' for k, v in links_gerados.items()])
        
        html_content = f"""<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <title>Dossiê OSINT - {alvo}</title>
    <style>
        body {{ background-color: #0a0f1d; color: #c9d1d9; font-family: 'Segoe UI', sans-serif; padding: 30px; }}
        .container {{ max-width: 900px; margin: auto; background: #111827; border: 1px solid #1f2937; border-radius: 8px; padding: 25px; box-shadow: 0 4px 15px rgba(0,0,0,0.5); }}
        h1 {{ color: #38bdf8; font-size: 24px; border-bottom: 2px solid #1f2937; padding-bottom: 10px; }}
        .meta-info {{ font-size: 13px; color: #9ca3af; margin-bottom: 20px; }}
        .section-title {{ color: #f43f5e; font-size: 16px; font-weight: bold; margin: 25px 0 10px 0; text-transform: uppercase; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; background: #0f172a; border-radius: 6px; overflow: hidden; }}
        td {{ padding: 12px 15px; border-bottom: 1px solid #1e293b; }}
        .label {{ color: #38bdf8; font-weight: bold; width: 30%; }}
        .value {{ color: #ffffff; }}
        .grid-links {{ display: grid; grid-template-columns: 1fr; gap: 12px; }}
        .link-card {{ background: #1e293b; padding: 12px 15px; border-radius: 6px; display: flex; justify-content: space-between; align-items: center; border-left: 4px solid #38bdf8; }}
        .link-title {{ font-weight: 600; color: #e2e8f0; }}
        .btn-link {{ background: #0284c7; color: #ffffff; text-decoration: none; padding: 6px 12px; border-radius: 4px; font-size: 12px; font-weight: bold; }}
        .footer {{ margin-top: 40px; text-align: center; font-size: 11px; color: #4b5563; border-top: 1px solid #1f2937; padding-top: 15px; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📑 Relatório de Inteligência Passiva (OSINT)</h1>
        <div class="meta-info">
            <strong>Operação:</strong> Triagem de Vetores Civis e Financeiros<br>
            <strong>Data/Hora de Emissão:</strong> {timestamp_humano}<br>
            <strong>Geração:</strong> Automação Willian OSINT Multi-Tool
        </div>
        <div class="section-title">🔍 I - Descrição Textual do Alvo e Qualificação</div>
        <table> {linhas_tabela} </table>
        <div class="section-title">🌐 II - Fontes de Cruzamento Vinculadas (Links Clicáveis)</div>
        <div class="grid-links"> {botoes_links} </div>
        <div class="footer"> Dossiê digital gerado localmente de forma automatizada por Willian de Oliveira. </div>
    </div>
</body>
</html>"""
        with open(filepath_html, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(f"\n{GREEN}[✅] DOSSIÊ HTML GERADO COM SUCESSO: {os.path.abspath(filepath_html)}{RESET}")
        
        # Dispara a geração simultânea do Word na mesma pasta
        self.gerar_backup_word(self.dirs["civil"], nome_base_arquivo, f"Relatório Técnico de Inteligência OSINT - {alvo}", dados_extra, links_gerados)
        webbrowser.open('file://' + os.path.realpath(filepath_html))

    def modulo_civil(self):
        print(f"\n{BLUE}[MÓDULO CIVIL]{RESET}")
        phone = input("Digite o número de Telefone (ex: 16999999999): ").strip()
        nome = input("Digite o Nome Completo ou Alcunha (ex: Nome completo): ").strip()
        if not nome:
            print(f"{RED}[!] Campo Nome é obrigatório para dorks.{RESET}")
            return
        
        nome_url = urllib.parse.quote(nome)
        ddd_info = self.obter_info_ddd(phone)
        
        links = {
            "Histórico de Processos (JusBrasil)": f"https://www.jusbrasil.com.br/busca?q={nome_url}",
            "Antecedentes e Publicações (Escavador)": f"https://www.escavador.com/busca?q={nome_url}",
            "Vínculos e Menções Web (Google Citações)": f"https://www.google.com/search?q=%22{nome_url}%22",
            "Menções Textuais em PDFs Públicos": f"https://www.google.com/search?q=%22{nome_url}%22+filetype:pdf",
            "Vínculos Profissionais (LinkedIn)": f"https://www.google.com/search?q=site:linkedin.com/in/+%22{nome_url}%22"
        }
        
        dados_extra = {
            "Telefone Alvo": phone,
            "Região Estimada (Análise DDD)": ddd_info,
            "Termo Nominativo": nome,
            "Status OPSEC": "100% Passivo (Sem rastro na API)"
        }
        self.criar_dossie_html_civil("CIVIL", nome, dados_extra, links)

    def modulo_infraestrutura(self):
        print(f"\n{BLUE}[MÓDULO INFRAESTRUTURA - ANÁLISE DE IP / DOMÍNIOS]{RESET}")
        alvo = input("Digite o IP ou Domínio Suspeito (ex: site-falso.com): ").strip()
        if not alvo: return
        
        links = {
            "Consulta Cadastral de IP/Domínio (Whois)": f"https://who.is/whois/{alvo}",
            "Reputação e Histórico de Ameaças (VirusTotal)": f"https://www.virustotal.com/gui/search/{alvo}",
            "Localização e ASN de Infraestrutura": f"https://bgp.he.net/dns/{alvo}",
            "Histórico de Prints e Alterações (Wayback Machine)": f"https://web.archive.org/web/*/{alvo}"
        }
        dados_extra = {
            "Vetor de Rede": alvo,
            "Categoria da Análise": "Infraestrutura Maliciosa / Phishing",
            "Auditoria Técnica": "Investigação Passiva DNS"
        }
        self.criar_dossie_html_civil("INFRA", alvo, dados_extra, links)

    def modulo_telegram_web(self):
        print(f"\n{BLUE}[MÓDULO TELEGRAM OSINT WEB (SEM API)]{RESET}")
        username = input("Digite o username do alvo (sem @ - ex: criminoso_tg): ").strip()
        if not username: return
        
        links = {
            "Visualização de Perfil (Telegram Web Público)": f"https://t.me/{username}",
            "Histórico de Menções Globais no Google": f"https://www.google.com/search?q=%22{username}%22",
            "Vínculos do Username em outras redes (Insta/TT)": f"https://www.google.com/search?q=%22{username}%22+-site:t.me"
        }
        dados_extra = {
            "Username Coletado": f"@{username}",
            "Estratégia Utilizada": "Dorking de vazamento de ID/Identidade",
            "Modo": "Stealth Completo"
        }
        self.criar_dossie_html_civil("TG_WEB", username, dados_extra, links)

    # =========================================================================
    # ESTEIRA 2: OSINT JURÍDICO E PATRIMONIAL
    # =========================================================================
    def criar_dossie_html_juridico(self, tipo_doc, documento, dados_extra, links_gerados):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        timestamp_humano = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        
        nome_base_arquivo = f"RELATORIO_{tipo_doc}_{documento}_{timestamp}"
        filename_html = f"{nome_base_arquivo}.html"
        filepath_html = os.path.join(self.dirs["juridico"], filename_html)
        
        linhas_tabela = "".join([f'<tr><td class="label">{k}</td><td class="value">{v}</td></tr>' for k, v in dados_extra.items()])
        botoes_links = "".join([f'<div class="link-card"><span class="link-title">{k}</span><a href="{v}" target="_blank" class="btn-link">Auditar Fonte Oficial</a></div>' for k, v in links_gerados.items()])
        
        html_content = f"""<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <style>
        body {{ background-color: #f8fafc; color: #1e293b; font-family: 'Times New Roman', serif; padding: 40px; }}
        .documento {{ max-width: 850px; margin: auto; background: #ffffff; border: 2px solid #334155; padding: 40px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); }}
        .timbre {{ text-align: center; font-weight: bold; font-size: 18px; border-bottom: 3px double #334155; padding-bottom: 15px; margin-bottom: 30px; }}
        h1 {{ font-size: 20px; text-align: center; color: #0f172a; }}
        .meta-info {{ font-size: 14px; background: #f1f5f9; padding: 15px; border-left: 5px solid #475569; margin-bottom: 30px; }}
        .secao-titulo {{ font-size: 15px; font-weight: bold; background: #334155; padding: 6px 12px; margin-top: 30px; color: white; text-transform: uppercase; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 25px; }}
        th, td {{ padding: 10px 12px; border: 1px solid #cbd5e1; font-size: 14px; }}
        .label {{ background: #f8fafc; font-weight: bold; width: 35%; }}
        .grid-links {{ display: flex; flex-direction: column; gap: 10px; margin-top: 15px; }}
        .link-card {{ background: #ffffff; padding: 10px 15px; border: 1px solid #e2e8f0; display: flex; justify-content: space-between; align-items: center; }}
        .btn-link {{ background: #1e293b; color: white; text-decoration: none; padding: 5px 10px; font-size: 11px; font-weight: bold; }}
    </style>
</head>
<body>
    <div class="documento">
        <div class="timbre">Relatório Técnico de Extração de Evidências Cadastrais e Patrimoniais</div>
        <h1>RELATÓRIO DE INTELIGÊNCIA FORENSE DIGITAL</h1>
        <div class="meta-info">
            <strong>Natureza:</strong> Levantamento Patrimonial / Triagem Cadastral Jurídica<br>
            <strong>Data/Hora de Emissão:</strong> {timestamp_humano}<br>
            <strong>Alvo da Consulta:</strong> {documento} ({tipo_doc})
        </div>
        <div class="secao-titulo">I - Qualificação Inicial Informada</div>
        <table> {linhas_tabela} </table>
        <div class="secao-titulo">II - Fontes Oficiais e Verificação de Vínculos (Links Clicáveis)</div>
        <div class="grid-links"> {botoes_links} </div>
        <div class="secao-titulo">III - Metodologia Forense</div>
        <p style="font-size: 13px; text-align: justify;">Os links estruturados utilizam herança de dorks avançadas em indexes públicos e cartorários para localização de fraudes à execução civil ou ocultações em holding.</p>
    </div>
</body>
</html>"""
        with open(filepath_html, "w", encoding="utf-8") as f:
            f.write(html_content)
        print(f"\n{GREEN}[✅] RELATÓRIO JURÍDICO HTML GERADO: {os.path.abspath(filepath_html)}{RESET}")
        
        # Dispara o Word correspondente para a esteira jurídica
        self.gerar_backup_word(self.dirs["juridico"], nome_base_arquivo, f"Dossiê Forense Patrimonial - {documento}", dados_extra, links_gerados)
        webbrowser.open('file://' + os.path.realpath(filepath_html))

    def analisar_cpf(self):
        print(f"\n{BLUE}[PESQUISA AVANÇADA POR CPF]{RESET}")
        cpf_raw = input("Digite o CPF do Alvo: ").strip()
        nome = input("Digite o Nome Completo do Executado/Réu: ").strip()
        cpf = self.limpar_documento(cpf_raw)
        
        nome_url = urllib.parse.quote(nome)
        links = {
            "Pesquisa de Processos Unificada (JusBrasil)": f"https://www.jusbrasil.com.br/busca?q={nome_url}",
            "Consulta Cadastral Corporativa (Transparência CC)": f"https://transparencia.cc/buscar/socios/?q={nome_url}",
            "Rastreamento de Diários Oficiais (Escavador)": f"https://www.escavador.com/busca?q={nome_url}",
            "Histórico Cadastral de Endereços em Editais": f"https://www.google.com/search?q=%22{nome_url}%22+AND+(%22edital%22+OR+%22citado%22+OR+%22intima%C3%A7%C3%A3o%22)"
        }
        dados = {"CPF Alvo": cpf, "Nome Completo": nome, "Tipo": "Pessoa Física"}
        self.criar_dossie_html_juridico("CPF", cpf, dados, links)

    def analisar_cnpj(self):
        print(f"\n{BLUE}[PESQUISA AVANÇADA POR CNPJ]{RESET}")
        cnpj_raw = input("Digite o CNPJ da Empresa: ").strip()
        cnpj = self.limpar_documento(cnpj_raw)
        
        links = {
            "Consulta de CNPJ Gratuita (CNPJ.biz)": f"https://cnpj.biz/{cnpj}",
            "Dados de Quadro Societário (Casa dos Dados)": f"https://casadosdados.com.br/solucao/cnpj/{cnpj}",
            "Rastreamento de Processos Corporativos": f"https://www.jusbrasil.com.br/busca?q={cnpj}"
        }
        dados = {"CNPJ Alvo": cnpj, "Tipo de Entrada": "Pessoa Jurídica"}
        self.criar_dossie_html_juridico("CNPJ", cnpj, dados, links)

    def executar_investigacao_judicial(self):
        print(f"\n{BLUE}[ESTEIRA OSINT: LITÍGIOS, DIVÓRCIOS & BENS]{RESET}")
        nome_alvo = input("[>] Nome Completo do Alvo para dorks judiciais: ").strip()
        if not nome_alvo: return
        
        nome_url = requests.utils.quote(nome_alvo)
        timestamp = datetime.now().strftime("%H%M%S")
        nome_base = f"RELATORIO_JUDICIAL_{nome_alvo.replace(' ', '_').upper()}_{timestamp}"
        filepath_html = os.path.join(self.dirs["juridico"], f"{nome_base}.html")
        
        dorks = {
            "Ocultação Patrimonial": f"https://www.google.com/search?q=%22{nome_url}%22+AND+(%22ocultacao+de+bens%22+OR+%22blindagem+patrimonial%22+OR+%22laranja%22+OR+%22holding+familiar%22)",
            "Partilha e Divórcio": f"https://www.google.com/search?q=%22{nome_url}%22+AND+(%22partilha+de+bens%22+OR+%22divorcio+litigioso%22)",
            "Ostentação / Padrão de Vida (Redes Sociais)": f"https://www.google.com/search?q=%22{nome_url}%22+AND+(%22viagem%22+OR+%22compras%22+OR+%22instagram%22+OR+%22facebook%22)"
        }
        
        html = f"""<html><body style="background:#000; color:#00ff00; font-family:monospace; padding:20px;">
        <div style="border:2px solid #00ff00; padding:20px;">
            <h1>⚖️ PAINEL FORENSE DE SINAIS DE RIQUEZA E LITÍGIOS</h1>
            <p><strong>Alvo:</strong> {nome_alvo}</p><hr>"""
        for k, v in dorks.items():
            html += f'<p><strong>{k}:</strong><br><a href="{v}" target="_blank" style="color:#00ff00;">{v}</a></p>'
        html += "</div></body></html>"
        
        with open(filepath_html, "w", encoding="utf-8") as f:
            f.write(html)
        print(f"{GREEN}[✅] Relatório de Blindagem salvo em: {filepath_html}{RESET}")
        
        # Word correspondente
        dados_doc = {"Alvo da Análise Judicial": nome_alvo, "Escopo": "Litígios & Sinais de Riqueza"}
        self.gerar_backup_word(self.dirs["juridico"], nome_base, f"Painel Forense de Sinais de Riqueza - {nome_alvo}", dados_doc, dorks)
        webbrowser.open('file://' + os.path.realpath(filepath_html))

    # =========================================================================
    # ESTEIRA 3: AUDITORIA E INTELIGÊNCIA FINANCEIRA
    # =========================================================================
    def modulo_financeiro_avancado(self):
        print(f"\n{BLUE}[MÓDULO DE INTELIGÊNCIA FINANCEIRA (CVM & BANCO CENTRAL)]{RESET}")
        alvo = input("Digite o Nome Completo ou CNPJ para Auditoria Financeira: ").strip()
        if not alvo: return
        
        alvo_url = urllib.parse.quote(alvo)
        tipo_entrada = "CNPJ" if any(char.isdigit() for char in alvo) else "Nome Completo/Alcunha"
        
        fontes = {
            "Consulta de Participantes CVM": f"https://dados.cvm.gov.br/pt-br/search?q={alvo_url}",
            "Acesso Direto ao Registrato (Banco Central)": "https://www.bcb.gov.br/meubc/registrato"
        }
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        alvo_limpo = re.sub(r'[^a-zA-Z0-9]', '_', alvo)
        nome_base = f"AUDITORIA_FINANCEIRA_{alvo_limpo}_{timestamp}"
        filepath_md = os.path.join(self.dirs["financeiro"], f"{nome_base}.md")
        
        md_content = f"""# 📑 Relatório de Auditoria e Inteligência Financeira
**Termo/Alvo Pesquisado:** {alvo}
**Tipo de Entrada:** {tipo_entrada}
**Data da Extração:** {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}
---
## 🔍 1. Fontes Oficiais Estruturadas
"""
        for nome_fonte, url in fontes.items():
            md_content += f"* **{nome_fonte}**: {url}\n"
            webbrowser.open_new_tab(url)
            
        with open(filepath_md, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"{GREEN}[✅] Relatório Markdown (.md) consolidado em: {os.path.abspath(filepath_md)}{RESET}")
        
        # Word correspondente
        dados_doc = {"Alvo Financeiro": alvo, "Vetor": tipo_entrada}
        self.gerar_backup_word(self.dirs["financeiro"], nome_base, f"Auditoria de Risco e Inteligência Financeira - {alvo}", dados_doc, fontes)

    # =========================================================================
    # ESTEIRA 4: CRUZAMENTO DE VETORES E CRIPTO
    # =========================================================================
    def modulo_cruzamento_centralizado(self):
        print(f"\n{BLUE}[SISTEMA DE INTERSECÇÃO CENTRALIZADA DE DADOS]{RESET}")
        print("[1] Módulo Investigativo de Chave PIX (Telefone/Email)")
        print("[2] Módulo de Rastreamento de Endereços Blockchain / Crypto")
        opcao = input("[>] Selecione o vetor de cruzamento: ").strip()
        
        alvo = input("[>] Digite o dado bruto (Chave PIX ou Endereço Wallet): ").strip()
        if not alvo: return
        
        links = {}
        tipo = ""
        extra = {}
        
        if opcao == '1':
            tipo = "PIX"
            nome_real = input("[?] Nome presumido do titular da chave PIX (opcional): ").strip()
            links = {
                "Buscar Processos Judiciais do Titular": f"https://www.jusbrasil.com.br/busca?q={urllib.parse.quote(nome_real)}",
                "Trabalho e Cargo (LinkedIn)": f"https://www.google.com/search?q=site:linkedin.com+%22{urllib.parse.quote(nome_real)}%22"
            }
            extra = {"Chave PIX Analisada": alvo, "Titular Presumido": nome_real if nome_real else "Não Fornecido"}
        elif opcao == '2':
            tipo = "CRIPTO"
            links = {
                "Rastrear Movimentações na Blockchain (Blockchair)": f"https://blockchair.com/search?q={alvo}",
                "Verificar Saldo e Tokens (Etherscan)": f"https://etherscan.io/address/{alvo}",
                "Consultar Blacklists Globais de Fraude/Scam": f"https://www.google.com/search?q=%22{alvo}%22+scam+fraud"
            }
            extra = {"Wallet Blockchain": alvo, "Rede": "Varredura Multichain Passiva"}
        else:
            return

        timestamp = datetime.now().strftime("%H%M%S")
        nome_base = f"DOSSIE_{tipo}_{timestamp}"
        filepath_html = os.path.join(self.dirs["centralizado"], f"{nome_base}.html")
        
        html = f"""<!DOCTYPE html><html lang="pt-br"><head><meta charset="UTF-8">
        <style>
            body {{ background: #000; color: #00ff41; font-family: 'Segoe UI', monospace; padding: 25px; }}
            .card {{ border: 2px solid #00ff41; padding: 20px; background: #050505; box-shadow: 0 0 15px #00ff41; }}
            h1 {{ color: #fff; border-bottom: 2px solid #ff0055; }}
            a {{ color: #00d4ff; text-decoration: none; display: block; margin: 8px 0; }}
        </style></head><body><div class="card">
        <h1>RELATÓRIO DE CRUZAMENTO DE INTELIGÊNCIA: {tipo}</h1>
        <p><strong>Dado de Entrada:</strong> {alvo}</p>"""
        for k, v in links.items():
            html += f'<a href="{v}" target="_blank">>> {k}</a>'
        html += "</div></body></html>"
        
        with open(filepath_html, "w", encoding="utf-8") as f:
            f.write(html)
        print(f"{GREEN}[✅] Dossiê de cruzamento persistido em: {filepath_html}{RESET}")
        
        # Word correspondente
        self.gerar_backup_word(self.dirs["centralizado"], nome_base, f"Cruzamento Centralizado de Inteligência Financeira ({tipo})", extra, links)
        webbrowser.open('file://' + os.path.realpath(filepath_html))

    # =========================================================================
    # ESTEIRA 5: TELEGRAM VIA API
    # =========================================================================
    async def _run_telegram_investigation(self, alvo):
        from telethon import TelegramClient
        print(f"\n{YELLOW}[*] Conectando à API do Telegram...{RESET}")
        client = TelegramClient('sessao_detetive', self.tg_api_id, self.tg_api_hash)
        await client.start()
        
        try:
            entidade = await client.get_entity(alvo)
            dados_reais = {
                "ID Numérico": entidade.id,
                "Primeiro Nome": entidade.first_name,
                "Sobrenome": entidade.last_name if entidade.last_name else "N/A",
                "Username": f"@{entidade.username}" if entidade.username else "Sem Username",
                "Telefone Oculto": entidade.phone if entidade.phone else "Privado (Ocultado por Privacidade)",
                "É Bot?": "Sim" if entidade.bot else "Não",
                "Bio/Descrição": "Extraível via busca profunda de logs"
            }
            
            nome_base = f"EVIDENCIA_TG_{dados_reais['ID Numérico']}"
            filepath_html = os.path.join(self.dirs["centralizado"], f"{nome_base}.html")
            
            html = f"""<html><body style='background:#000;color:#0f0;font-family:monospace;padding:30px;'>
            <h1>DOSSIÊ TELEGRAM: {dados_reais['Username']}</h1><hr>"""
            for k, v in dados_reais.items():
                html += f"<p><strong>{k}:</strong> {v}</p>"
            html += "</body></html>"
            
            with open(filepath_html, "w", encoding="utf-8") as f:
                f.write(html)
            print(f"{GREEN}[+] Alvo Identificado! ID Numérico Forense: {entidade.id}{RESET}")
            print(f"{GREEN}[✅] Evidência Digital em HTML salva em: {filepath_html}{RESET}")
            
            # Word correspondente para o Telegram
            self.gerar_backup_word(self.dirs["centralizado"], nome_base, f"Evidência Forense Digital - Telegram ID {entidade.id}", dados_reais, {})
            webbrowser.open('file://' + os.path.realpath(filepath_html))
        except Exception as e:
            print(f"{RED}[-] Erro crítico ao localizar alvo na API Telegram: {e}{RESET}")
        finally:
            await client.disconnect()

    def modulo_telegram_deep_api(self):
        print(f"\n{BLUE}[MÓDULO DE INVESTIGAÇÃO DE ALVOS NO TELEGRAM VIA API]{RESET}")
        alvo = input("Digite o @username ou o link t.me/ do alvo: ").strip()
        if not alvo: return
        
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            
        loop.run_until_complete(self._run_telegram_investigation(alvo))


# --- GERENCIADOR CENTRAL DO MENU PRINCIPAL (FLOW CONTROL) ---
def menu_supremo():
    engine = WillianOsintSupremo()
    
    while True:
        engine.limpar_tela()
        print(BANNER_PRINCIPAL)
        print(f"{BOLD}Selecione a Esteira ou Vetor Investigativo de Entrada:{RESET}")
        print(f" {GREEN}[1]{RESET} Módulo Civil Multi-Tool (Telefone + Menções Nominativas) [HTML+WORD]")
        print(f" {GREEN}[2]{RESET} Módulo de Infraestrutura Cibernética (IP / Phishing / Domínios) [HTML+WORD]")
        print(f" {GREEN}[3]{RESET} Telegram OSINT Web (Busca Rápida de Usernames Sem Uso de API) [HTML+WORD]")
        print(f" {GREEN}[4]{RESET} Pesquisa Avançada Jurídica por CPF (Pessoa Física / Executados) [HTML+WORD]")
        print(f" {GREEN}[5]{RESET} Pesquisa Avançada Jurídica por CNPJ (Empresas / Grupos Econômicos) [HTML+WORD]")
        print(f" {GREEN}[6]{RESET} Esteira de Litígios, Casamentos e Ocultação de Sinais de Riqueza [HTML+WORD]")
        print(f" {GREEN}[7]{RESET} Inteligência Financeira Avançada (Consultas CVM & Banco Central) [MD+WORD]")
        print(f" {GREEN}[8]{RESET} Cruzamento Centralizado de Vetores (Chaves PIX & Wallets Crypto) [HTML+WORD]")
        print(f" {GREEN}[9]{RESET} Rastreamento de Alvos e Extração de ID Telegram {RED}(Requer API){RESET} [HTML+WORD]")
        print(f" {RED}[0]{RESET} Encerrar Todos os Sistemas de Inteligência")
        print(f"{CYAN}==========================================================================={RESET}")
        
        opcao = input(f"{BOLD}Willian_Supremo_Console > {RESET}").strip()
        
        if opcao == '1':
            engine.modulo_civil()
        elif opcao == '2':
            engine.modulo_infraestrutura()
        elif opcao == '3':
            engine.modulo_telegram_web()
        elif opcao == '4':
            engine.analisar_cpf()
        elif opcao == '5':
            engine.analisar_cnpj()
        elif opcao == '6':
            engine.executar_investigacao_judicial()
        elif opcao == '7':
            engine.modulo_financeiro_avancado()
        elif opcao == '8':
            engine.modulo_cruzamento_centralizado()
        elif opcao == '9':
            engine.modulo_telegram_deep_api()
        elif opcao == '0':
            print(f"\n{YELLOW}[*] Desconectando barramentos. Relatórios armazenados localmente com segurança. Fim de Operação.{RESET}\n")
            break
        else:
            print(f"\n{RED}[!] Comando inválido. Escolha uma opção de 0 a 9.{RESET}")
        
        input(f"\n{CYAN}Pressione [Enter] para retornar ao Menu Principal...{RESET}")


if __name__ == "__main__":
    try:
        menu_supremo()
    except KeyboardInterrupt:
        print(f"\n\n{RED}[!] Interrupção forçada pelo operador. Deslogando com OPSEC ativa.{RESET}\n")
        sys.exit(0)
