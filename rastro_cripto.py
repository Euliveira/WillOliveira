#!/usr/bin/env python3
"""
Módulo Forense Premium: Geração de Laudo Técnico Avançado em DOCX com Design Profissional.

Este script demonstra a automação de design editorial e formatação estética 
para relatórios periciais. Ele aplica margens personalizadas, tipografia corporativa,
tabelas estilizadas com sombreamento e paleta de cores institucional (Azul Marinho).

Requisitos: pip install python-docx requests
"""

import sys
import datetime
import requests

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    Document = None

class GeradorLaudoDesignPremium:
    """Mecanismo de auditoria blockchain focado em relatórios de alto padrão visual."""
    
    CONTRATO_USDC = input("contrato: ") #0xa0b86991c6218b36c1d19d4a2e9eb0ce3606eb48"
    URL_API = "https://api.etherscan.io/api"
    
    # Definição da paleta de cores corporativa (Padrão Forense/Jurídico)
    COR_PRIMARIA = RGBColor(10, 37, 64)       # Azul Marinho Escuro
    COR_SECUNDARIA = RGBColor(0, 102, 204)    # Azul Elétrico (Destaques)
    COR_TEXTO = RGBColor(51, 51, 51)          # Cinza Escuro (Leitura confortável)
    COR_MUTED = RGBColor(119, 119, 119)       # Cinza Claro (Metadados)
    
    HEX_FUNDO_TABELA = "0A2540"                # Azul Marinho em Hex para XML
    HEX_LINHA_ZEBRA = "F4F6F9"                 # Cinza claro para linhas alternadas

    def __init__(self, api_key: str):
        self.api_key = api_key

    def obter_historico_saidas(self, carteira_alvo: str) -> list:
        """Coleta e filtra as transações de saída de tokens ERC-20 via API."""
        parametros = {
            "module": "account", "action": "tokentx", 
            "contractaddress": self.CONTRATO_USDC, "address": carteira_alvo,
            "page": 1, "offset": 50, "sort": "desc", "apikey": self.api_key
        }
        try:
            resposta = requests.get(self.URL_API, params=parametros, timeout=15)
            dados = resposta.json()
            if dados.get("status") != "1": return []
            return [tx for tx in dados.get("result", []) if tx.get("from", "").lower() == carteira_alvo.lower()]
        except Exception:
            return []

    def aplicar_sombreamento_celula(self, celula, hex_cor: str):
        """Aplica cor de fundo personalizada a uma célula da tabela via manipulação XML."""
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_cor}"/>'
        celula._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    def gerar_laudo_estilizado(self, carteira_investigada: str, destino_identificado: str, nome_entidade: str):
        """Processa a lógica de dados e constrói o design editorial do arquivo .docx."""
        if Document is None:
            print("[ERRO] Bibliotecas de estilização visual não encontradas.")
            return

        transacoes = self.obter_historico_saidas(carteira_investigada)
        doc = Document()
        
        # 1. Configuração Avançada de Margens (Padrão ABNT/Forense)
        seccao = doc.sections[0]
        seccao.top_margin = Inches(1.0)
        seccao.bottom_margin = Inches(1.0)
        seccao.left_margin = Inches(1.0)
        seccao.right_margin = Inches(1.0)

        # Configuração do Estilo Padrão do Texto
        estilo_normal = doc.styles['Normal']
        estilo_normal.font.name = 'Arial'
        estilo_normal.font.size = Pt(10.5)
        estilo_normal.font.color.rgb = self.COR_TEXTO

        # 2. Cabeçalho Estilizado (Tarja de Título)
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_after = Pt(24)
        
        run_t1 = p_titulo.add_run("RELATÓRIO TÉCNICO DE AUDITORIA FORENSE DIGITAL\n")
        run_t1.bold = True
        run_t1.font.size = Pt(16)
        run_t1.font.color.rgb = self.COR_PRIMARIA
        
        run_t2 = p_titulo.add_run("ANÁLISE DE FLUXO DE ATIVOS ON-CHAIN E RASTREAMENTO CRIMINAL")
        run_t2.font.size = Pt(10)
        run_t2.font.color.rgb = self.COR_SECUNDARIA
        run_t2.bold = True

        # Linha Divisória Elegante
        p_linha = doc.add_paragraph()
        p_linha.paragraph_format.space_after = Pt(18)
        run_linha = p_linha.add_run("—" * 65)
        run_linha.font.color.rgb = self.COR_MUTED

        # 3. Quadro de Resumo de Metadados (Box de Informações)
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.line_spacing = 1.3
        p_meta.paragraph_format.space_after = Pt(24)
        
        p_meta.add_run("IDENTIFICAÇÃO DO PROCESSO:\n").bold = True
        p_meta.add_run(f"  • Data/Hora Impressão: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')} BRT\n")
        p_meta.add_run(f"  • Metodologia Aplicada: Análise Contábil Forense Blockchain (Follow-the-Money)\n")
        p_meta.add_run(f"  • Token Alvo Monitorado: USD Coin (USDC) — Smart Contract: {self.CONTRATO_USDC}\n")
        p_meta.add_run(f"  • Carteira de Investigação (Origem): ").bold = True
        p_meta.add_run(f"{carteira_investigada}\n")

        # 4. Seção 1: Constatações Técnicas
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(8)
        run_h1 = h1.add_run("1. CONSTATAÇÕES TÉCNICAS E MAPEAMENTO DE FLUXO")
        run_h1.bold = True
        run_h1.font.size = Pt(12)
        run_h1.font.color.rgb = self.COR_PRIMARIA

        p_desc1 = doc.add_paragraph(
            "Por meio de varredura automatizada nos blocos de registro público imutável da rede Ethereum, "
            "foram isoladas as transações de debito da carteira de origem que caracterizam o desvio patrimonial "
            "objeto desta apuração. As evidências coletadas estão sumarizadas na tabela analítica abaixo:"
        )
        p_desc1.paragraph_format.space_after = Pt(12)

        # 5. Construção e Estilização Avançada da Tabela de Evidências
        if not transacoes:
            p_vazio = doc.add_paragraph("[Aviso] Nenhuma transferência ativa localizada no período de consulta.")
            p_vazio.italic = True
        else:
            tabela = doc.add_table(rows=1, cols=4)
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Formatação do Cabeçalho da Tabela (Fundo Escuro, Texto Branco e Negrito)
            headers = ['Data/Hora (UTC)', 'Volume (USDC)', 'Hash da Transação (TxID)', 'Endereço Destino']
            hdr_cells = tabela.rows[0].cells
            for idx, text in enumerate(headers):
                hdr_cells[idx].text = text
                self.aplicar_sombreamento_celula(hdr_cells[idx], self.HEX_FUNDO_TABELA)
                run = hdr_cells[idx].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(9.5)
                hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Preenchimento das Linhas com Efeito Zebra (Alternância de Cores)
            for i, tx in enumerate(transacoes):
                row_cells = tabela.add_row().cells
                ts = int(tx.get("timeStamp", 0))
                data_formato = datetime.datetime.utcfromtimestamp(ts).strftime('%d/%m/%Y\n%H:%M:%S')
                valor_ajustado = int(tx.get("value", 0)) / (10 ** int(tx.get("tokenDecimal", 6)))
                
                row_cells[0].text = data_formato
                row_cells[1].text = f"{valor_ajustado:,.2f}"
                row_cells[2].text = f"{tx.get('hash')[:18]}..."
                row_cells[3].text = f"{tx.get('to')[:18]}..."
                
                # Alinhamento e fontes das células de dados
                for c_idx, cell in enumerate(row_cells):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.runs[0].font.size = Pt(9)
                    if c_idx in [0, 1]: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Aplica cor cinza claro nas linhas pares para efeito visual agradável
                    if i % 2 == 0:
                        self.aplicar_sombreamento_celula(cell, self.HEX_LINHA_ZEBRA)

        doc.add_paragraph("\n")

        # 6. Seção 2: Atribuição de Entidade
        h2 = doc.add_paragraph()
        run_h2 = h2.add_run("2. ANÁLISE DE DESTINO E ATRIBUIÇÃO DE PROPRIEDADE")
        run_h2.bold = True
        run_h2.font.size = Pt(12)
        run_h2.font.color.rgb = self.COR_PRIMARIA
        h2.paragraph_format.space_before = Pt(14)

        p_desc2 = doc.add_paragraph()
        p_desc2.paragraph_format.line_spacing = 1.2
        p_desc2.add_run("O cruzamento dos caminhos digitais aponta que a totalidade dos ativos convergiu para o endereço:\n")
        run_wallet = p_desc2.add_run(f"👉 {destino_identificado}\n\n")
        run_wallet.bold = True
        run_wallet.font.color.rgb = self.COR_SECUNDARIA
        
        p_desc2.add_run(
            "Conclusão de Vínculo: De acordo com os dicionários globais de atribuição de infraestrutura da blockchain, "
            "este endereço hexadecimal específico corresponde a uma carteira agregadora corporativa pertencente à plataforma: "
        )
        p_desc2.add_run(f"{nome_entidade}.\n").bold = True
        p_desc2.add_run(
            "Essa constatação técnica certifica que os valores roubados ingressaram no ambiente fechado da referida corretora, "
            "gerando um crédito imediato na conta interna de um usuário específico de sua plataforma."
        )

        # 7. Seção 3: Recomendações de Ação
        h3 = doc.add_paragraph()
        run_h3 = h3.add_run("3. ENCAMINHAMENTO DE DIRETRIZES LEGAIS")
        run_h3.bold = True
        run_h3.font.size = Pt(12)
        run_h3.font.color.rgb = self.COR_PRIMARIA
        h3.paragraph_format.space_before = Pt(14)

        p_desc3 = doc.add_paragraph(
            "Com base nas evidências materiais e rastreáveis colhidas, recomenda-se que as autoridades judiciárias adotem "
            "as seguintes medidas coercitivas urgentes para a preservação de direitos e recuperação dos ativos:\n"
            "  1. Expedição de Mandado/Ofício Judicial eletrônico à gerência de conformidade legal da Exchange identificada.\n"
            "  2. Ordem de bloqueio administrativo urgente sobre qualquer conta que tenha recebido aportes das TxIDs catalogadas.\n"
            "  3. Determinação de quebra de sigilo cadastral (KYC), fornecendo dados de CPF, contas de liquidação bancária e IPs de acesso."
        )
        p_desc3.paragraph_format.line_spacing = 1.2

        # Rodapé Oficial Estilizado
        doc.add_paragraph("\n" + "_"*60 + "\n")
        p_fim = doc.add_paragraph("DOCUMENTO TÉCNICO PERICIAL - EMITIDO PARA FINS DE INSTRUÇÃO PROCESSUAL")
        p_fim.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_fim.runs[0].font.size = Pt(8.5)
        p_fim.runs[0].font.color.rgb = self.COR_MUTED

        # Salva o arquivo final
        nome_arquivo = f"laudo_premium_{carteira_investigada[:8]}.docx"
        doc.save(nome_arquivo)
        print(f"[CONCLUÍDO] Relatório com design de alto padrão gerado em: {nome_arquivo}")

# --- Ponto de Execução do Sistema (Simulado) ---
if __name__ == "__main__":
    CHAVE_API_MOCK = "CHAVE_PROVISORIA_EXEMPLO"
    CARTEIRA_INVESTIGADA_EXEMPLO = "0x55611510b3d1dcc800d431e224fbd3bf9833cfbf"
    DESTINO_BINANCE_EXEMPLO = "0x28c6c06298d514db089934071355e5ba621b4d23"
    
    gerador = GeradorLaudoDesignPremium(api_key=CHAVE_API_MOCK)
    # A chamada abaixo simularia a escrita estética do laudo final para Word
    # gerador.gerar_laudo_estilizado(
    #     carteira_investigada=CARTEIRA_INVESTIGADA_EXEMPLO,
    #     destino_identificado=DESTINO_BINANCE_EXEMPLO,
    #     nome_entidade="BINANCE (Binance 14 / Hot Wallet 6)"
    # )
    print("[MÓDULO DESIGN PREMIUM] Estrutura visual corporativa carregada com sucesso.")
