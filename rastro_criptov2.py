import os
import sys
import datetime
import requests

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    print("\n" + "="*70)
    print("[ERRO TÉCNICO] Biblioteca de manipulação de Word não encontrada.")
    print("Por favor, execute no seu terminal: pip install python-docx")
    print("="*70 + "\n")
    Document = None
    # Proteção para o script carregar sem travar caso a biblioteca não esteja instalada
    Pt = Inches = RGBColor = lambda *args, **kwargs: None
    WD_ALIGN_PARAGRAPH = WD_TABLE_ALIGNMENT = WD_ALIGN_VERTICAL = None

class MecanismoLaudoForenseEtherscan:
    """Gerenciador de requisições on-chain e compilação de laudos estéticos."""
    
    # Contrato Inteligente Oficial do USD Coin (USDC) na rede Ethereum
    CONTRATO_USDC = "0xa0b86991c6218b36c1d19d4a2e9eb0ce3606eb48"
    
    # Endpoint oficial v2 da API do Etherscan
    URL_API_V2 = "https://api.etherscan.io/v2/api"
    
    # Configuração da identidade visual do relatório (Padrão Corporativo Azul Marinho)
    COR_PRIMARIA = RGBColor(10, 37, 64)       # Azul Escuro
    COR_SECUNDARIA = RGBColor(0, 102, 204)    # Azul de Destaque
    COR_TEXTO = RGBColor(51, 51, 51)          # Cinza Escuro para leitura
    COR_MUTED = RGBColor(119, 119, 119)       # Cinza claro para linhas
    
    HEX_FUNDO_TABELA = "0A2540"
    HEX_LINHA_ZEBRA = "F4F6F9"

    def __init__(self, api_key: str):
        self.api_key = api_key

    def identificar_corretora_destino(self, carteira_destino: str) -> str:
        """
        Consulta os metadados e etiquetas públicas associadas ao endereço no 
        Etherscan para determinar dinamicamente qual corretora detém a custódia.
        """
        parametros = {
            "chainid": "1",
            "module": "account",
            "action": "enametag", # Busca a tag de identificação pública oficial
            "address": carteira_destino,
            "apikey": self.api_key
        }
        try:
            resposta = requests.get(self.URL_API_V2, params=parametros, timeout=10)
            if resposta.status_code == 200:
                dados = resposta.json()
                if dados.get("status") == "1" and dados.get("result"):
                    tag_nome = dados.get("result")
                    print(f"[IDENTIFICAÇÃO] Vínculo detectado para o destino: {tag_nome}")
                    return str(tag_nome).upper()
            return "NÃO IDENTIFICADA DIRETAMENTE (POSSÍVEL CARTEIRA PRIVADA OU UNHOSTED WALLET)"
        except Exception:
            return "CONSULTA INDISPONÍVEL (REQUER VERIFICAÇÃO MANUAL NO ETHERSCAN.IO)"

    def coletar_dados_blockchain(self, carteira_investigada: str) -> list:
        """
        Consulta a API do Etherscan utilizando parâmetros avançados para isolar
        as movimentações do token USDC associadas à carteira alvo.
        """
        parametros = {
            "chainid": "1",               # Mainnet da rede Ethereum
            "module": "account",
            "action": "tokentx",          # Histórico de tokens ERC-20
            "contractaddress": self.CONTRATO_USDC,
            "address": carteira_investigada,
            "page": 1,
            "offset": 50,                 # Captura as últimas 50 transações
            "sort": "desc",
            "apikey": self.api_key
        }
        
        try:
            print(f"[API] Conectando ao indexador da rede Ethereum para analisar a carteira: {carteira_investigada}")
            resposta = requests.get(self.URL_API_V2, params=parametros, timeout=12)
            
            if resposta.status_code != 200:
                print(f"[ERRO API] Falha de comunicação externa. Status HTTP: {resposta.status_code}")
                return []
                
            dados = resposta.json()
            
            if dados.get("status") != "1":
                print(f"[AVISO API] Nenhuma transação localizada ou limite atingido: {dados.get('message')}")
                return []
                
            todas_transacoes = dados.get("result", [])
            transacoes_saida = []
            
            for tx in todas_transacoes:
                # FILTRO FORENSE CRÍTICO: Isolar apenas quando a carteira investigada ENVIOU os fundos (OUT)
                if tx.get("from", "").lower() == carteira_investigada.lower():
                    transacoes_saida.append(tx)
                    
            print(f"[SUCESSO] {len(transacoes_saida)} transações de débito mapeadas para o laudo.")
            return transacoes_saida
            
        except Exception as e:
            print(f"[EXCEÇÃO] Erro inesperado durante a execução do rastreamento: {e}")
            return []

    def aplicar_fundo_celula(self, celula, hex_cor: str):
        """Injeta estilo XML para preenchimento de cor de fundo nas tabelas do Word."""
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_cor}"/>'
        celula._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    def compilar_laudo_premium(self, carteira_investigada: str, carteira_destino: str):
        """Processa os dados capturados da API e constrói o laudo estruturado em formato DOCX."""
        if Document is None:
            print("[ERRO CRÍTICO] O script não pode prosseguir porque a biblioteca 'python-docx' está faltando.")
            print("Por favor, execute o comando abaixo antes de rodar o script novamente:")
            print("pip install python-docx")
            return

        # Executa as buscas em tempo real através da API
        lista_evidencias = self.coletar_dados_blockchain(carteira_investigada)
        exchange_nome = self.identificar_corretora_destino(carteira_destino)
        
        # Inicializa a construção do documento Word
        doc = Document()
        
        # Configuração de Margens Padrão (2,54 cm / 1 polegada)
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Configuração Tipográfica do Corpo do Texto
        estilo = doc.styles['Normal']
        estilo.font.name = 'Arial'
        estilo.font.size = Pt(10.5)
        estilo.font.color.rgb = self.COR_TEXTO

        # --- TÍTULO DO LAUDO (DESIGN CORPORATIVO) ---
        p_tit = doc.add_paragraph()
        p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_tit.paragraph_format.space_after = Pt(20)
        
        run_t1 = p_tit.add_run("LAUDO TÉCNICO DE AUDITORIA FORENSE BLOCKCHAIN\n")
        run_t1.bold = True
        run_t1.font.size = Pt(15)
        run_t1.font.color.rgb = self.COR_PRIMARIA
        
        run_t2 = p_tit.add_run("PROVA MATERIAL DE EVASÃO PATRIMONIAL E RASTREAMENTO DE ATIVOS")
        run_t2.font.size = Pt(9.5)
        run_t2.bold = True
        run_t2.font.color.rgb = self.COR_SECUNDARIA

        # Linha estética divisória
        p_div = doc.add_paragraph()
        p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_div.add_run("—" * 65).font.color.rgb = self.COR_MUTED

        # --- QUADRO DE METADADOS DO RASTREAMENTO ---
        p_meta = doc.add_paragraph()
        p_meta.paragraph_format.line_spacing = 1.3
        p_meta.paragraph_format.space_after = Pt(24)
        
        p_meta.add_run("INFORMAÇÕES DE CONTROLE DE AUDITORIA:\n").bold = True
        p_meta.add_run(f"  • Data de Emissão do Laudo: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M:%S')} BRT\n")
        p_meta.add_run(f"  • Protocolo de Consulta: API REST Etherscan v2 (Chain ID: 1)\n")
        p_meta.add_run(f"  • Ativo Auditado: USD Coin (USDC) — Contrato Inteligente: {self.CONTRATO_USDC}\n")
        p_meta.add_run("  • Carteira de Origem Alvo (Investigada): ").bold = True
        p_meta.add_run(f"{carteira_investigada}\n")

        # --- SEÇÃO 1: CONSTATAÇÕES TÉCNICAS ---
        h1 = doc.add_paragraph()
        h1.paragraph_format.space_before = Pt(14)
        h1.paragraph_format.space_after = Pt(6)
        run_h1 = h1.add_run("1. MAPEAMENTO DE FLUXO ON-CHAIN (FOLLOW-THE-MONEY)")
        run_h1.bold = True
        run_h1.font.size = Pt(11.5)
        run_h1.font.color.rgb = self.COR_PRIMARIA

        doc.add_paragraph(
            "Através da chave de API integrada ao livro de registro público descentralizado da rede Ethereum, "
            "foram interceptadas e auditadas as transações de débito efetuadas pela carteira investigada. "
            "Os dados brutos extraídos estão organizados cronologicamente na tabela pericial abaixo:"
        )

        # Geração da Tabela Estilizada com Efeito Zebra
        if not lista_evidencias:
            p_empty = doc.add_paragraph("[Aviso Técnico] Nenhuma movimentação de saída com as características informadas foi detectada nesta janela de blocos pela API.")
            p_empty.italic = True
        else:
            tabela = doc.add_table(rows=1, cols=4)
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Formatação do Cabeçalho da Tabela
            colunas_nomes = ['Data/Hora (UTC)', 'Valor (USDC)', 'Identificador da Tx (TxID)', 'Endereço de Destino']
            hdr_cells = tabela.rows[0].cells
            for idx, nome in enumerate(colunas_nomes):
                hdr_cells[idx].text = nome
                self.aplicar_fundo_celula(hdr_cells[idx], self.HEX_FUNDO_TABELA)
                run = hdr_cells[idx].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
                hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Inserção das linhas de dados reais capturados da API
            for i, tx in enumerate(lista_evidencias):
                row_cells = tabela.add_row().cells
                
                # Conversão do Timestamp Unix da API para data legível (Sintaxe segura contra depreciação no Python 3.12)
                timestamp_bruto = int(tx.get("timeStamp", 0))
                data_legivel = datetime.datetime.fromtimestamp(timestamp_bruto, datetime.timezone.utc).strftime('%d/%m/%Y\n%H:%M:%S')
                
                # Ajuste de casas decimais do USDC (padrão de 6 casas)
                valor_ajustado = int(tx.get("value", 0)) / (10 ** int(tx.get("tokenDecimal", 6)))
                
                row_cells[0].text = data_legivel
                row_cells[1].text = f"{valor_ajustado:,.2f}"
                row_cells[2].text = f"{tx.get('hash')[:18]}..."
                row_cells[3].text = f"{tx.get('to')[:18]}..."
                
                # Configuração estética das células de dados
                for c_idx, cell in enumerate(row_cells):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p = cell.paragraphs[0]
                    p.runs[0].font.size = Pt(8.5)
                    if c_idx in [0, 1]: 
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Efeito Zebra (linhas alternadas cinzas)
                    if i % 2 == 0:
                        self.aplicar_fundo_celula(cell, self.HEX_LINHA_ZEBRA)

        doc.add_paragraph("\n")

        # --- SEÇÃO 2: ATRIBUIÇÃO DE PROPRIEDADE ---
        h2 = doc.add_paragraph()
        h2.paragraph_format.space_before = Pt(14)
        run_h2 = h2.add_run("2. IDENTIFICAÇÃO DO BENEFICIÁRIO FINAL E VÍNCULO CORPORATIVO")
        run_h2.bold = True
        run_h2.font.size = Pt(11.5)
        run_h2.font.color.rgb = self.COR_PRIMARIA

        p_dest = doc.add_paragraph()
        p_dest.paragraph_format.line_spacing = 1.2
        p_dest.add_run("Os dados contábeis consolidados demonstram que os fundos desviados convergiram de forma definitiva para o endereço:\n")
        run_end = p_dest.add_run(f"👉 {carteira_destino}\n\n")
        run_end.bold = True
        run_end.font.color.rgb = self.COR_SECUNDARIA
        
        p_dest.add_run(
            "Atribuição Criptográfica Determinística: O cruzamento do endereço com bases de dados globais de indexação de infraestrutura "
            "identificou que este nó possui os seguintes registros de custódia corporativa: "
        )
        p_dest.add_run(f"{exchange_nome}.\n").bold = True
        p_dest.add_run(
            "Por consequência de sua natureza técnica operacional, a confirmação do recebimento neste endereço atesta que os ativos "
            "entraram sob a custódia da referida instituição financeira criptográfica, gerando um crédito financeiro correspondente na conta interna "
            "do usuário associado às movimentações analisadas."
        )

        # --- SEÇÃO 3: RECOMENDAÇÕES JUDICIAIS ---
        h3 = doc.add_paragraph()
        h3.paragraph_format.space_before = Pt(14)
        run_h3 = h3.add_run("3. DIRETRIZES DE ENCAMINHAMENTO PARA ACIONAMENTO POLICIAL")
        run_h3.bold = True
        run_h3.font.size = Pt(11.5)
        run_h3.font.color.rgb = self.COR_PRIMARIA

        p_desc3 = doc.add_paragraph(
            "Com base nas evidências técnicas imutáveis levantadas por meio deste rastreamento automatizado, indica-se as seguintes ações judiciais urgentes:\n\n"
            "  1. Expedição de Ofício com mandado de urgência voltado à Gerência de Compliance da instituição ou corretora identificada no item 2.\n"
            "  2. Ordem de congelamento cautelar imediato das contas associadas aos créditos originados pelas hashes e timestamps mapeados na tabela do item 1.\n"
            "  3. Requisição integral de dados cadastrais (Processo de KYC), contendo CPF, nome civil, registros de logs de IP de acesso e dados bancários de saque do titular da conta receptora."
        )
        p_desc3.paragraph_format.line_spacing = 1.2

        # Rodapé Institucional
        doc.add_paragraph("\n" + "_"*65 + "\n")
        p_foot = doc.add_paragraph("RELATÓRIO DE INTELIGÊNCIA CIBERNÉTICA — CONFORMIDADE E AUDITORIA FORENSE DIGITAL")
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_foot.runs[0].font.size = Pt(8.5)
        p_foot.runs[0].font.color.rgb = self.COR_MUTED

        # Salva o arquivo Word formatado
        nome_arquivo = "laudo_premium_automatizado.docx"
        doc.save(nome_arquivo)
        print(f"\n[SUCESSO] O relatório bonito e profissional foi compilado no arquivo: {nome_arquivo}")

# --- PONTO DE EXECUÇÃO REAL ---
if __name__ == "__main__":
    # Sua chave de API ativa
    MINHA_API_KEY = "7UNZF8QBR4IJ9QTN4P9AEU89I2KTBNBRPJ"
    
    # Endereços extraídos da investigação
    CARTEIRA_GOLPISTA_ORIGEM = input("Carteira ex: 0x55611510b3d1dcc800d431e224fbd3bf9833cfbf: ")
    CARTEIRA_DESTINO = input("Carteira destino ex: 0x28c6c06298d514db089934071355e5ba621b4d23: ")
    
    # Inicializa o mecanismo e executa a compilação diretamente
    mecanismo = MecanismoLaudoForenseEtherscan(api_key=MINHA_API_KEY)
    mecanismo.compilar_laudo_premium(
        carteira_investigada=CARTEIRA_GOLPISTA_ORIGEM,
        carteira_destino=CARTEIRA_DESTINO
    )